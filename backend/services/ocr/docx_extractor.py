"""
DOCX text extraction using python-docx.
Treats entire document as page 1 (DOCX has no native page boundaries).
"""
import logging
from typing import List

from docx import Document as DocxDocument

from models.bidder import PageContent

logger = logging.getLogger(__name__)


def extract_docx(file_path: str) -> List[PageContent]:
    """
    Extract text from a DOCX file.

    Extracts all paragraphs in order, then all tables row by row, cell by cell.
    Returns a single PageContent (page 1) since DOCX has no page boundaries.

    Args:
        file_path: Path to the DOCX file

    Returns:
        List with single PageContent containing the full document text
    """
    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        logger.error(f"Failed to open DOCX {file_path}: {e}")
        raise

    text_parts = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                text_parts.append(" | ".join(row_texts))

    full_text = "\n".join(text_parts)
    word_count = len(full_text.split())

    logger.info(f"Extracted DOCX {file_path}: {word_count} words")

    return [PageContent(
        page_num=1,
        raw_text=full_text,
        word_count=word_count,
        extraction_method="docx",
        needs_documentai=False,
    )]
